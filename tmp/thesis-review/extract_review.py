from __future__ import annotations

import json
import re
import zipfile
from collections import Counter
from pathlib import Path

import pdfplumber
from docx import Document
from lxml import etree


ROOT = Path(r"D:\Programming\Thesis\LLMLogAnalyzer")
DOCS = ROOT / "documents"
OUT = ROOT / "tmp" / "thesis-review"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def docx_report(path: Path, output: Path) -> None:
    doc = Document(path)
    paragraphs = []
    styles = Counter()
    run_fonts = Counter()
    run_sizes = Counter()
    for idx, para in enumerate(doc.paragraphs):
        text = clean(para.text)
        style = para.style.name if para.style else ""
        styles[style] += 1
        if text:
            paragraphs.append(
                {
                    "index": idx,
                    "style": style,
                    "alignment": str(para.alignment),
                    "text": text,
                }
            )
        for run in para.runs:
            if clean(run.text):
                run_fonts[run.font.name or "(inherited)"] += len(clean(run.text))
                run_sizes[str(run.font.size.pt) if run.font.size else "(inherited)"] += len(clean(run.text))

    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([clean(cell.text) for cell in row.cells])
        tables.append({"index": ti, "rows": rows})

    sections = []
    for si, section in enumerate(doc.sections):
        sections.append(
            {
                "index": si,
                "start_type": str(section.start_type),
                "page_width": section.page_width,
                "page_height": section.page_height,
                "top_margin": section.top_margin,
                "bottom_margin": section.bottom_margin,
                "left_margin": section.left_margin,
                "right_margin": section.right_margin,
                "header_distance": section.header_distance,
                "footer_distance": section.footer_distance,
            }
        )

    with zipfile.ZipFile(path) as zf:
        document_xml = zf.read("word/document.xml")
        root = etree.fromstring(document_xml)
        xml_text = " ".join(clean(t) for t in root.xpath(".//w:t/text()", namespaces=NS) if clean(t))
        tracked_insertions = len(root.xpath(".//w:ins", namespaces=NS))
        tracked_deletions = len(root.xpath(".//w:del", namespaces=NS))
        comment_ranges = len(root.xpath(".//w:commentRangeStart", namespaces=NS))
        fields = [clean(x) for x in root.xpath(".//w:instrText/text()", namespaces=NS)]
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
        comments = []
        if "word/comments.xml" in zf.namelist():
            croot = etree.fromstring(zf.read("word/comments.xml"))
            for c in croot.xpath(".//w:comment", namespaces=NS):
                comments.append(clean(" ".join(c.xpath(".//w:t/text()", namespaces=NS))))

    cp = doc.core_properties
    words = re.findall(r"\S+", xml_text)
    report = {
        "path": str(path),
        "core_properties": {
            "title": cp.title,
            "subject": cp.subject,
            "author": cp.author,
            "last_modified_by": cp.last_modified_by,
            "created": cp.created.isoformat() if cp.created else None,
            "modified": cp.modified.isoformat() if cp.modified else None,
        },
        "counts": {
            "paragraphs_total": len(doc.paragraphs),
            "paragraphs_nonempty": len(paragraphs),
            "tables": len(tables),
            "sections": len(doc.sections),
            "images": len(media),
            "words_approx": len(words),
            "tracked_insertions": tracked_insertions,
            "tracked_deletions": tracked_deletions,
            "comment_ranges": comment_ranges,
            "comments": len(comments),
        },
        "style_counts": styles.most_common(),
        "run_font_char_counts": run_fonts.most_common(),
        "run_size_char_counts": run_sizes.most_common(),
        "sections": sections,
        "paragraphs": paragraphs,
        "tables": tables,
        "fields": fields,
        "comments": comments,
        "all_xml_text": xml_text,
    }
    output.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def pdf_report(path: Path, output: Path) -> None:
    pages = []
    with pdfplumber.open(path) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=2, y_tolerance=2, layout=False) or ""
            pages.append(
                {
                    "page": idx,
                    "width": page.width,
                    "height": page.height,
                    "images": len(page.images),
                    "tables_detected": len(page.find_tables()),
                    "text": text,
                }
            )
    output.write_text(
        json.dumps({"path": str(path), "page_count": len(pages), "pages": pages}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def main() -> None:
    docx_report(DOCS / "پایان نامه ی مسعود دباغی.docx", OUT / "main_docx.json")
    docx_report(DOCS / "فرم پروپوزال کارشناسی ارشد مسعود دباغی.docx", OUT / "proposal_docx.json")
    pdf_report(OUT / "main" / "main.pdf", OUT / "main_pdf.json")
    pdf_report(OUT / "proposal" / "proposal.pdf", OUT / "proposal_pdf.json")
    pdf_report(DOCS / "فرمت پایان نامه و رساله.pdf", OUT / "guideline_pdf.json")
    pdf_report(DOCS / "پایان-نامه-محمد نوخیز.pdf", OUT / "sample_pdf.json")


if __name__ == "__main__":
    main()
