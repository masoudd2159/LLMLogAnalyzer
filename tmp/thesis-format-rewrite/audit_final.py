import hashlib
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


ROOT = Path(r"D:\Programming\Thesis\LLMLogAnalyzer")
WORKING = ROOT / "documents" / "پایان نامه ی مسعود دباغی - نسخه کاری.docx"
ORIGINAL = ROOT / "documents" / "پایان نامه ی مسعود دباغی.docx"
EXPECTED_ORIGINAL_SHA = "4E3C640E49BCD23D4044B79ED0C94BFF3BDB6C9F781ABC3CF2D2117E870D23D8"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


doc = Document(WORKING)
with zipfile.ZipFile(WORKING) as archive:
    bad_member = archive.testzip()
    document_xml = etree.fromstring(archive.read("word/document.xml"))
    footnotes_xml = etree.fromstring(archive.read("word/footnotes.xml"))

refs = [
    int(node.get(qn("w:id")))
    for node in document_xml.xpath(".//w:footnoteReference", namespaces=NS)
]
defs = [
    int(node.get(qn("w:id")))
    for node in footnotes_xml.xpath("./w:footnote", namespaces=NS)
    if int(node.get(qn("w:id"))) >= 1
]
instructions = [
    "".join(node.itertext())
    for node in document_xml.xpath(".//w:instrText", namespaces=NS)
]
all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
all_text += "\n" + "\n".join(
    cell.text for table in doc.tables for row in table.rows for cell in row.cells
)

persian_body = []
missing_bidi = []


def has_effective_bidi(paragraph):
    if paragraph._p.pPr is not None and paragraph._p.pPr.find(qn("w:bidi")) is not None:
        return True
    style = paragraph.style
    visited = set()
    while style is not None and style.style_id not in visited:
        visited.add(style.style_id)
        ppr = style.element.pPr
        if ppr is not None and ppr.find(qn("w:bidi")) is not None:
            return True
        style = style.base_style
    return False


for index, paragraph in enumerate(doc.paragraphs[141:394], start=141):
    text = paragraph.text.strip()
    if text and re.search(r"[\u0600-\u06FF]", text):
        persian_body.append(index)
        if not has_effective_bidi(paragraph):
            missing_bidi.append(index)

table_bidi_missing = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if re.search(r"[\u0600-\u06FF]", paragraph.text):
                    if not has_effective_bidi(paragraph):
                        table_bidi_missing += 1

styles = {}
for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Caption1", "Footnote Text"):
    try:
        style = doc.styles[name]
    except KeyError:
        styles[name] = "missing"
        continue
    fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    styles[name] = {
        "size_pt": style.font.size.pt if style.font.size else None,
        "bold": style.font.bold,
        "cs": fonts.get(qn("w:cs")),
        "eastAsia": fonts.get(qn("w:eastAsia")),
    }

section_report = []
for section in doc.sections:
    section_report.append(
        {
            "page_cm": (round(section.page_width.cm, 2), round(section.page_height.cm, 2)),
            "margins_cm": (
                round(section.top_margin.cm, 2),
                round(section.right_margin.cm, 2),
                round(section.bottom_margin.cm, 2),
                round(section.left_margin.cm, 2),
            ),
            "footer_cm": round(section.footer_distance.cm, 2),
        }
    )

references = [
    paragraph.text
    for paragraph in doc.paragraphs
    if re.match(r"^\[\d+\]", paragraph.text.strip())
]
table_one_ids = {
    int(match.group(1))
    for row in doc.tables[1].rows[1:]
    for match in [re.search(r"\b(\d{1,2})\b", row.cells[0].text)]
    if match
}

print(f"zip_ok={bad_member is None}")
print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} rows={[len(t.rows) for t in doc.tables]}")
print(f"footnote_refs={len(refs)} footnote_defs={len(defs)} ids_match={refs == defs == list(range(1, 38))}")
print(f"tc_fields={sum(code.strip().startswith('TC ') for code in instructions)} toc_fields={sum(code.strip().startswith('TOC ') for code in instructions)}")
print(f"unresolved_markers={len(re.findall(r'\[\[FN\d+\]\]', all_text))}")
print(f"field_errors={sum(term in all_text for term in ('Error! Bookmark not defined.', 'No table of contents entries found.'))}")
print(f"persian_body_paragraphs={len(persian_body)} missing_bidi={missing_bidi}")
print(f"table_bidi_missing={table_bidi_missing}")
print(f"styles={styles}")
print(f"sections={section_report}")
print(f"references={len(references)}")
print(f"article_table_ids={sorted(table_one_ids)}")
print(f"original_sha_ok={sha256(ORIGINAL) == EXPECTED_ORIGINAL_SHA}")
