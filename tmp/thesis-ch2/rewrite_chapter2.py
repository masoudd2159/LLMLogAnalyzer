from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Programming\Thesis\LLMLogAnalyzer")
DOCX = ROOT / "documents" / "پایان نامه ی مسعود دباغی - نسخه کاری.docx"
BACKUP = ROOT / "tmp" / "thesis-ch2" / "before-chapter2-rewrite.docx"


def set_bidi(paragraph, enabled=True):
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if enabled and bidi is None:
        ppr.append(OxmlElement("w:bidi"))
    elif not enabled and bidi is not None:
        ppr.remove(bidi)


def set_run_fonts(run, size=14, bold=None, persian="B Nazanin", latin="Times New Roman"):
    run.font.name = latin
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._r.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), persian)
    fonts.set(qn("w:cs"), persian)


def keep(paragraph, next_paragraph=False, together=True):
    ppr = paragraph._p.get_or_add_pPr()
    for tag, enabled in (("w:keepNext", next_paragraph), ("w:keepLines", together)):
        old = ppr.find(qn(tag))
        if enabled and old is None:
            ppr.append(OxmlElement(tag))
        elif not enabled and old is not None:
            ppr.remove(old)


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    set_bidi(p, True)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    keep(p, next_paragraph=True, together=True)
    return p


def add_body(doc, text, *, center=False, keep_next=False, ltr=False, size=14):
    p = doc.add_paragraph(style="Normal")
    r = p.add_run(text)
    set_run_fonts(r, size=size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_bidi(p, not ltr)
    keep(p, next_paragraph=keep_next, together=True)
    return p


def add_caption(doc, text, identifier):
    caption_style = "Caption" if "Caption" in [s.name for s in doc.styles] else "Caption1"
    p = doc.add_paragraph(style=caption_style)
    r = p.add_run(text)
    set_run_fonts(r, size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_bidi(p, True)
    keep(p, next_paragraph=(identifier == "T"), together=True)
    add_hidden_tc_field(p, text, identifier)
    return p


def add_hidden_tc_field(paragraph, entry, identifier="T"):
    begin = paragraph.add_run()
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin._r.append(begin_char)
    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f'TC "{entry}" \\f {identifier} \\l "1"'
    instr._r.append(instr_text)
    end = paragraph.add_run()
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end._r.append(end_char)


def add_complex_field(paragraph, instruction, placeholder=""):
    paragraph.clear()
    begin = paragraph.add_run()
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin._r.append(begin_char)
    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    instr._r.append(instr_text)
    separate = paragraph.add_run()
    separate_char = OxmlElement("w:fldChar")
    separate_char.set(qn("w:fldCharType"), "separate")
    separate._r.append(separate_char)
    if placeholder:
        result = paragraph.add_run(placeholder)
        set_run_fonts(result, size=12)
    end = paragraph.add_run()
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end._r.append(end_char)
    set_bidi(paragraph, True)


def paragraph_after(doc, paragraph):
    ps = doc.paragraphs
    idx = next(i for i, item in enumerate(ps) if item._p is paragraph._p)
    return ps[idx + 1]


def fresh_field_paragraph_after(title_paragraph):
    """Remove old generated TOC-result paragraphs and return a clean field paragraph."""
    parent = title_paragraph._parent
    node = title_paragraph._p.getnext()
    while node is not None and node.tag == qn("w:p"):
        candidate = Paragraph(node, parent)
        if not candidate.style.name.lower().startswith("toc"):
            break
        nxt = node.getnext()
        node.getparent().remove(node)
        node = nxt
    if node is not None and node.tag == qn("w:p"):
        candidate = Paragraph(node, parent)
        if not candidate.text.strip() and candidate.style.name == "Normal":
            return candidate
    new_p = OxmlElement("w:p")
    title_paragraph._p.addnext(new_p)
    return Paragraph(new_p, parent)


def replace_front_fields(doc):
    titles = {
        "فهرست مطالب": 'TOC \\o "1-3" \\h \\z \\u',
        "فهرست جدول ها": "TOC \\f T \\h \\z",
        "فهرست شکل ها": "TOC \\f F \\h \\z",
    }
    for title, field in titles.items():
        p = next((x for x in doc.paragraphs if x.text.strip() == title), None)
        if p is not None:
            body = fresh_field_paragraph_after(p)
            add_complex_field(body, field, "فهرست پس از به‌روزرسانی فیلدهای Word نمایش داده می‌شود.")


def remove_after_chapter_heading(doc):
    chapter = next(
        p for p in doc.paragraphs
        if p.style.name == "Heading 1"
        and p.text.strip().startswith("فصل دوم: مروری بر ادبیات و پیشینه تحقیق")
    )
    node = chapter._p.getnext()
    body = doc._body._element
    while node is not None and node.tag != qn("w:sectPr"):
        nxt = node.getnext()
        body.remove(node)
        node = nxt
    return chapter


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    trpr.append(marker)


def set_table_rtl(table):
    tblpr = table._tbl.tblPr
    bidi = tblpr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tblpr.append(bidi)
    bidi.set(qn("w:val"), "1")


def no_split(row):
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:cantSplit")) is None:
        trpr.append(OxmlElement("w:cantSplit"))


def format_cell(cell, text, *, size=9, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_run_fonts(r, size=size, bold=bold)
    p.alignment = align
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_bidi(p, True)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    if fill:
        shade(cell, fill)


def add_data_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_rtl(table)
    repeat_header(table.rows[0])
    no_split(table.rows[0])
    for idx, (header, width) in enumerate(zip(headers, widths)):
        table.rows[0].cells[idx].width = width
        format_cell(table.rows[0].cells[idx], header, size=9, bold=True, fill="D9EAF7")
    for row_data in rows:
        row = table.add_row()
        for idx, (value, width) in enumerate(zip(row_data, widths)):
            row.cells[idx].width = width
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in (1, 3, 4) else WD_ALIGN_PARAGRAPH.JUSTIFY
            format_cell(row.cells[idx], value, size=font_size, align=align)
    return table


def add_relationship_figure(doc):
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_rtl(table)
    widths = [Cm(4.0)] * 4
    for row in table.rows:
        no_split(row)
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    top = table.cell(0, 0).merge(table.cell(0, 3))
    format_cell(top, "ادبیات تشخیص ناهنجاری لاگ", size=12, bold=True, fill="4472C4")
    for run in top.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)

    categories = [
        ("روش‌های ترتیبی", "DeepLog؛ مدل‌سازی توالی رخداد"),
        ("مدل‌های پیش‌آموزش‌دیده", "LogBERT، BERT-Log و LogFiT"),
        ("LLM با آموزش/تنظیم", "LogGPT، LogLLM، LLM-LADE و PEFT"),
        ("پرامپت، بازیابی و ترکیب", "LogPrompt، RAGLog، AdaptiveLog و FlexLog"),
    ]
    for idx, (title, examples) in enumerate(categories):
        format_cell(table.cell(1, idx), title, size=10.5, bold=True, fill="D9EAF7")
        format_cell(table.cell(2, idx), examples, size=9.5, fill="F3F7FB")

    bridge = table.cell(3, 0).merge(table.cell(3, 3))
    format_cell(
        bridge,
        "همگرایی مؤلفه‌های موردنیاز: معناشناسی LLM + نرمال‌سازی قالب + کنترل تصمیم + کاهش فراخوانی مدل",
        size=10.5,
        bold=True,
        fill="FFF2CC",
    )

    labels = ["دامنه و ورودی", "دو مسیر تصمیم", "کارایی و قابلیت بازتولید", "استقرار و ارزیابی"]
    details = [
        "فقط BGL؛ تصمیم خط‌به‌خط؛ حذف برچسب از ورودی",
        "Hybrid: قاعده قطعی سپس LLM\nPrompt-only: تصمیم صرفاً در پرامپت",
        "کش سطح قالب نرمال‌شده؛ اعتبارسنجی JSON؛ نسخه‌بندی کلید کش",
        "Qwen2.5:7B-Instruct روی Ollama محلی؛ Accuracy، Precision، Recall، F1، زمان و INVALID",
    ]
    for idx in range(4):
        format_cell(table.cell(4, idx), labels[idx], size=10, bold=True, fill="E2F0D9")
        format_cell(table.cell(5, idx), details[idx], size=9.5, fill="F4FBF1")
    return table


def add_chapter_content(doc):
    add_heading(doc, "۲-۱ مقدمه", 2)
    add_body(doc, "لاگ‌های سیستمی ردپای رخدادهای زمان اجرا، وضعیت اجزا و پیام‌های خطا را ثبت می‌کنند و ازاین‌رو یکی از ورودی‌های اصلی پایش، عیب‌یابی و تشخیص ناهنجاری در سامانه‌های نرم‌افزاری به‌شمار می‌روند. افزایش حجم و تنوع لاگ‌ها، بررسی دستی را دشوار کرده و پژوهش‌ها را از قواعد و شمارش رخدادها به‌سوی مدل‌های ترتیبی، مدل‌های زبانی پیش‌آموزش‌دیده و در سال‌های اخیر مدل‌های زبانی بزرگ سوق داده است [1، 2، 22].")
    add_body(doc, "در روش‌های مبتنی بر یادگیری عمیق، توالی قالب‌های لاگ یا بازنمایی متنی پیام‌ها برای یادگیری الگوی رفتار عادی استفاده می‌شود. DeepLog با پیش‌بینی رخداد بعدی، LogBERT با دو وظیفه خودنظارتی و BERT-Log با تنظیم یک مدل BERT برای طبقه‌بندی، سه مسیر شاخص این تحول را نشان می‌دهند [1–3]. این روش‌ها توانایی مدل‌سازی الگو را افزایش داده‌اند، اما معمولاً به آموزش اختصاصی، ساخت توالی، تجزیه لاگ یا داده برچسب‌خورده وابسته‌اند [2، 3، 22].")
    add_body(doc, "مدل‌های زبانی بزرگ مسیر دیگری را فراهم کرده‌اند: انجام تشخیص از طریق دستور متنی، نمونه‌های درون‌متنی، بازیابی شواهد، تنظیم کامل یا تنظیم پارامترکارا. مطالعات LogGPT، LogPrompt، RAGLog، LogLLM، LLM-LADE، AdaptiveLog و FlexLog نشان می‌دهند که نقش مدل زبانی می‌تواند از طبقه‌بند مستقیم تا مفسر، تولیدکننده داده، جزء کمکی یا تصمیم‌گیر پشتیبان تغییر کند [4–8، 10، 12، 15]. بنابراین مقایسه پژوهش‌ها تنها با نام مدل کافی نیست و باید نوع آموزش، واحد تحلیل، شیوه آماده‌سازی ورودی، محل استقرار و سازوکار کنترل هزینه نیز در نظر گرفته شود [22، 24].")
    add_body(doc, "فصل حاضر مبانی موردنیاز برای پژوهش تشخیص ناهنجاری در لاگ‌های BGL با Qwen2.5:7B-Instruct و Ollama را ارائه می‌کند. سپس پیشینه در پنج دسته مرور و با جدول و شکل مقایسه می‌شود. در پایان، شکاف تحقیقاتی با اتکا به مجموعه منابع بررسی‌شده و جایگاه دو روش مورد مطالعه، یعنی Hybrid و Prompt-only، تبیین می‌شود. واژه Hybrid در این پژوهش به «قاعده قطعی با فراخوانی LLM در موارد باقیمانده» اشاره دارد و با اصطلاح Hybrid Prompting در مطالعه پارک و همکاران، که ترکیبی از شواهد سطح رخداد و توالی را در پرامپت قرار می‌دهد، یکسان نیست [9].")

    add_heading(doc, "۲-۲ مبانی نظری و مفاهیم موردنیاز", 2)
    add_heading(doc, "۲-۲-۱ لاگ سیستمی و تشخیص ناهنجاری", 3)
    add_body(doc, "هر پیام لاگ معمولاً از سرآیندی مانند زمان، سطح شدت و مؤلفه تولیدکننده و یک بدنه متنی تشکیل می‌شود. در بدنه، بخش‌های ثابت بیانگر رخداد برنامه و بخش‌های متغیر شامل شناسه‌ها، نشانی‌ها، مسیرها، مقادیر عددی و سایر پارامترهای زمان اجرا هستند. پژوهش‌های LogPPT و LogParser-LLM فرایند تبدیل پیام خام به قالب ثابت و پارامترهای متغیر را «تجزیه لاگ» تعریف می‌کنند و نشان می‌دهند که کیفیت این مرحله بر تحلیل‌های پایین‌دستی، از جمله تشخیص ناهنجاری، اثر می‌گذارد [18، 19].")
    add_body(doc, "در ادبیات، ناهنجاری به انحراف مشاهده از الگوی عادی آموخته‌شده یا تعریف‌شده اطلاق می‌شود. این انحراف می‌تواند در یک پیام منفرد، فراوانی رخدادها، ترتیب قالب‌ها یا معنای یک توالی ظاهر شود. DeepLog و LogBERT بر انحراف توالی از رفتار عادی تمرکز دارند؛ در مقابل ADALog هر پیام ساخت‌نیافته را مستقل تحلیل و احتمال بازسازی توکن‌ها را به امتیاز سطح پیام تبدیل می‌کند [1، 2، 17]. این تفاوت، واحد تحلیل و در نتیجه نحوه ساخت نمونه آزمایشی را تعیین می‌کند.")

    add_heading(doc, "۲-۲-۲ انواع ناهنجاری و واحد تحلیل", 3)
    add_body(doc, "سه سطح رایج در پژوهش‌های لاگ قابل تفکیک است: ناهنجاری رخداد یا خط، ناهنجاری توالی و ناهنجاری کمّی. در سطح خط، هر پیام یک نمونه تصمیم است و برچسب به همان پیام تعلق دارد. در سطح توالی، پیام‌ها براساس شناسه نشست، گره یا پنجره زمانی گروه‌بندی می‌شوند و وجود یک رخداد غیرعادی می‌تواند برچسب کل توالی را تعیین کند. در ناهنجاری کمّی، تغییر غیرعادی فراوانی قالب‌ها در پنجره موردنظر بررسی می‌شود [1–3، 22].")
    add_body(doc, "واحد تحلیل بر قابلیت مقایسه نتایج اثر مستقیم دارد. برای نمونه، BERT-Log در BGL از پنجره لغزان مبتنی بر شناسه گره، اندازه پنجره و گام استفاده کرده است؛ Hybrid Prompting توالی‌های یک‌دقیقه‌ای می‌سازد؛ اما ADALog تصمیم را برای یک پیام مستقل گزارش می‌کند [3، 9، 17]. ازاین‌رو اعداد دقت یا F1 مطالعاتی که خط، نشست یا پنجره زمانی را نمونه می‌گیرند بدون همسان‌سازی پروتکل ارزیابی قابل مقایسه مستقیم نیستند.")

    add_heading(doc, "۲-۲-۳ مجموعه‌داده BGL", 3)
    add_body(doc, "مجموعه‌داده Blue Gene/L یا BGL از لاگ‌های ابررایانه Blue Gene/L در آزمایشگاه ملی لارنس لیورمور تهیه شده است. در این مجموعه ۴٬۷۴۷٬۹۶۳ پیام لاگ وجود دارد که ۳۴۸٬۴۶۰ پیام، معادل حدود ۷٫۳۴ درصد، با برچسب هشدار به‌عنوان ناهنجار مشخص شده‌اند [2، 6]. وجود برچسب در ابتدای هر خط امکان ارزیابی نظارت‌شده را فراهم می‌کند، اما آن برچسب باید پیش از ارسال متن به مدل حذف شود تا نشت پاسخ به ورودی رخ ندهد.")
    add_body(doc, "مطالعات مختلف BGL را با پروتکل‌های متفاوت به‌کار گرفته‌اند. LogBERT پنجره زمانی پنج‌دقیقه‌ای با میانگین طول ۵۶۲ رخداد ساخته است؛ BERT-Log از پنجره مبتنی بر گره استفاده کرده و Hybrid Prompting پنجره یک‌دقیقه‌ای و نمونه متوازن از هزار توالی عادی و هزار توالی ناهنجار را ارزیابی کرده است [2، 3، 9]. در پژوهش حاضر، مسئله در سطح خط تعریف می‌شود؛ بنابراین برچسب مرجع هر خط فقط برای ارزیابی نگه‌داری و از ورودی Qwen حذف می‌شود.")

    add_heading(doc, "۲-۲-۴ تجزیه، نرمال‌سازی و قالب لاگ", 3)
    add_body(doc, "هدف تجزیه لاگ جداکردن کلمات ثابت قالب از پارامترهای متغیر است. روش‌های سنتی از فراوانی، شباهت یا قواعد نحوی استفاده می‌کنند؛ LogPPT با پرامپت‌تیونینگ و تعداد کمی نمونه برچسب‌خورده، و LogParser-LLM با ترکیب درخت پیشوند و فراخوانی محدود LLM، معنا را نیز وارد فرایند استخراج قالب کرده‌اند [18، 19]. LogParser-LLM در ارزیابی LogPub با میانگین ۳٫۶ میلیون لاگ در هر مجموعه، به‌طور متوسط با ۲۷۲٫۵ فراخوانی LLM، F1 گروه‌بندی ۹۰٫۶ درصد و دقت تجزیه ۸۱٫۱ درصد گزارش کرده است [18].")
    add_body(doc, "نرمال‌سازی در این پژوهش یک تجزیه کامل مبتنی بر خوشه‌بندی نیست؛ بلکه تبدیل قطعی مقادیر پویا به جانشین‌های پایدار برای ساخت قالب نرمال‌شده است. شناسه گره، واحدها، نشانی IP، مقادیر هگزادسیمال، تاریخ و زمان، مسیرها و اعداد پویا ماسک می‌شوند. بااین‌حال صفر و غیرصفر با دو جانشین متمایز <ZERO> و <NON_ZERO> نگه داشته می‌شوند، زیرا برخی پیام‌های عملیاتی تفاوت معنایی خود را از همین وضعیت عددی می‌گیرند. این طراحی با اصل جداسازی قالب و پارامتر در مطالعات تجزیه لاگ سازگار است، اما قاعده جانشینی آن به‌طور مشخص برای خط لوله این پژوهش تعریف شده است [18، 19].")
    add_body(doc, "یک نتیجه عملی نرمال‌سازی، کاهش تنوع ظاهری پیام‌هایی است که یک رخداد واحد را با شناسه یا عدد متفاوت بیان می‌کنند. در نتیجه، تصمیم سطح قالب می‌تواند میان خطوط هم‌معنا بازاستفاده شود. در این حالت دقت نرمال‌سازی اهمیت دارد: ماسک بیش‌ازحد ممکن است نشانه معنایی را حذف کند و ماسک کم نیز تعداد قالب‌ها و فراخوانی مدل را افزایش می‌دهد. مطالعات LogPPT، LogParser-LLM و LogLAA نیز حساسیت تحلیل پایین‌دستی به کیفیت آماده‌سازی و تجزیه را گزارش کرده‌اند [18، 19، 21].")

    add_heading(doc, "۲-۲-۵ روش‌های قاعده‌محور، یادگیری ماشین و مدل‌های ترتیبی", 3)
    add_body(doc, "روش‌های قاعده‌محور الگوهای از پیش تعریف‌شده را با واژه کلیدی یا عبارت منظم تطبیق می‌دهند. مزیت آن‌ها تصمیم سریع و قابل ردیابی برای الگوهای شناخته‌شده است، ولی دامنه کشف به قواعد نوشته‌شده محدود می‌شود و نگه‌داری قواعد در برابر تغییر قالب‌ها هزینه‌بر است [2، 22]. روش‌های کلاسیک یادگیری ماشین، پس از استخراج بردار فراوانی یا ویژگی‌های آماری، از الگوریتم‌هایی مانند PCA، ماشین بردار پشتیبان تک‌کلاسه یا درخت تصمیم برای تشخیص انحراف استفاده می‌کنند [1، 10].")
    add_body(doc, "DeepLog لاگ را همانند توالی زبان مدل‌سازی و با LSTM رخداد بعدی را از رخدادهای قبلی پیش‌بینی می‌کند؛ رخدادی که در مجموعه پیش‌بینی‌های محتمل قرار نگیرد ناهنجار تلقی می‌شود. این روش آموزش افزایشی و یک جریان تشخیص و عیب‌یابی را نیز ارائه می‌کند [1]. روش‌های بعدی مانند LogBERT، وابستگی دوطرفه و وظایف خودنظارتی را جایگزین اتکای صرف به پیش‌بینی رخداد بعدی کرده‌اند [2]. این خانواده برای ناهنجاری‌های ترتیبی مناسب است، اما به تعریف توالی و آموزش روی داده سامانه وابستگی دارد.")

    add_heading(doc, "۲-۲-۶ مدل‌های زبانی پیش‌آموزش‌دیده و مدل‌های زبانی بزرگ", 3)
    add_body(doc, "مدل‌های پیش‌آموزش‌دیده مانند BERT بازنمایی زمینه‌مند واژه‌ها را از متن فرا می‌گیرند و سپس برای وظیفه خاص تنظیم می‌شوند. LogBERT دو وظیفه «پیش‌بینی کلید ماسک‌شده» و «کمینه‌سازی حجم ابرکره» را بر توالی‌های عادی به‌کار می‌گیرد [2]. BERT-Log مدل BERT و یک لایه تمام‌متصل را برای طبقه‌بندی تنظیم کرده و در پروتکل خود F1 برابر ۹۹٫۴ درصد را برای BGL گزارش کرده است [3]. LogFiT نیز با مدل‌سازی توکن ماسک‌شده، فقط از لاگ‌های عادی برای یادگیری استفاده می‌کند [13].")
    add_body(doc, "مدل زبانی بزرگ معمولاً یک مدل Transformer با ظرفیت زیاد و پیش‌آموزی گسترده است که می‌تواند دستور متنی را دنبال و خروجی تولید کند. در حوزه تشخیص ناهنجاری، این مدل ممکن است مستقیماً نقش آشکارساز داشته باشد یا برای تولید داده، تبیین و پشتیبانی از مدل کوچک‌تر استفاده شود؛ این دو نقش در رده‌بندی مرور Xu و Ding به‌ترتیب «LLM برای تشخیص» و «LLM برای تولید» نام‌گذاری شده‌اند [24]. مرور نظام‌مند De la Cruz Cabello و همکاران نیز روش‌های لاگ را براساس آماده‌سازی، بازنمایی، مدل و شیوه ارزیابی دسته‌بندی می‌کند [22].")

    add_heading(doc, "۲-۲-۷ مدل Qwen2.5:7B-Instruct", 3)
    add_body(doc, "Qwen2.5 خانواده‌ای از مدل‌های متن‌باز با نسخه‌های پایه و دستورمحور در اندازه‌های ۰٫۵ تا ۷۲ میلیارد پارامتر است. گزارش فنی این خانواده، پیش‌آموزی بر ۱۸ تریلیون توکن و پس‌آموزی با بیش از یک میلیون نمونه تنظیم دستوری و مراحل تقویت یادگیری را شرح می‌دهد. نسخه 7B در این گزارش حدود ۶٫۵ میلیارد پارامتر غیرتعبیه‌ای دارد و نسخه Instruct برای پیروی از دستور و تولید پاسخ ساخت‌یافته بهینه شده است [25].")
    add_body(doc, "در این پژوهش Qwen2.5:7B-Instruct به‌صورت محلی و بدون تنظیم دقیق برای طبقه‌بندی دودویی هر قالب نرمال‌شده استفاده می‌شود. نبود تنظیم اختصاصی باعث می‌شود تفاوت دو روش آزمایش‌شده به معماری تصمیم و متن پرامپت مربوط باشد، نه تغییر وزن‌های مدل. انتخاب یک مدل و یک نسخه پرامپت ثابت برای هر دو روش، شرط لازم برای نسبت‌دادن اختلاف نتایج به سازوکار Hybrid یا Prompt-only است.")

    add_heading(doc, "۲-۲-۸ مهندسی پرامپت و خروجی ساخت‌یافته", 3)
    add_body(doc, "پرامپت مجموعه دستورها، زمینه دامنه، تعریف برچسب‌ها و قالب خروجی است که رفتار مدل را در زمان استنتاج هدایت می‌کند. LogGPT مبتنی بر ChatGPT اثر نوع پرامپت، اندازه پنجره و تزریق دانش انسانی را روی BGL و Spirit بررسی کرده است [6]. LogPrompt از راهبردهای چندمرحله‌ای و درخواست تبیین برای تجزیه و تشخیص ناهنجاری استفاده می‌کند و بدون آموزش درون‌دامنه روی چند مجموعه لاگ ارزیابی شده است [7]. OpsEval نیز حساسیت وظایف عملیات فناوری اطلاعات به پرامپت، زنجیره فکر و یادگیری درون‌متنی را در سطح معیار ارزیابی بررسی می‌کند [23].")
    add_body(doc, "برای طبقه‌بندی خودکار، خروجی باید از متن آزاد به قرارداد ماشین‌خوان تبدیل شود. قرارداد این پژوهش فقط یکی از دو شیء JSON یعنی {\"label\":\"0\"} برای عادی و {\"label\":\"1\"} برای ناهنجار را می‌پذیرد. پاسخ دارای متن اضافی، برچسب نامعتبر یا JSON ناقص با وضعیت INVALID ثبت می‌شود و وارد کش نمی‌گردد. مستندات رسمی Ollama امکان درخواست خروجی ساخت‌یافته با قالب یا طرح JSON و کاهش تصادفی‌بودن با دمای صفر را بیان می‌کند [26]. اعتبارسنجی مستقل پاسخ همچنان لازم است، زیرا قرارداد خروجی بخشی از پروتکل آزمایش و محاسبه نرخ پاسخ نامعتبر است.")

    add_heading(doc, "۲-۲-۹ استنتاج محلی با Ollama", 3)
    add_body(doc, "Ollama یک واسط محلی برای اجرای مدل و فراخوانی آن از طریق API فراهم می‌کند. API پیش‌فرض در نشانی localhost:11434/api در دسترس است و پاسخ‌های آن علاوه بر محتوای مدل، شاخص‌های زمان اجرا و شمار توکن را نیز گزارش می‌کنند [26]. استنتاج محلی مانع ارسال لاگ به خدمت خارجی می‌شود و کنترل نسخه مدل و پارامترهای اجرا را ممکن می‌سازد؛ مطالعه کاربرد واقعی De la Cruz Cabello و همکاران نیز استقرار محلی مدل LogBERT را برای محیط دارای محدودیت محرمانگی به‌کار گرفته است [27].")
    add_body(doc, "استنتاج محلی به معنی نبود هزینه محاسباتی نیست. زمان بارگذاری مدل، زمان ارزیابی پرامپت و زمان تولید پاسخ باید ثبت شوند و فراخوانی تکراری برای پیام‌های هم‌قالب کاهش یابد. مستندات Ollama فیلدهای total_duration، load_duration، prompt_eval_count، prompt_eval_duration، eval_count و eval_duration را برای اندازه‌گیری این اجزا معرفی می‌کند [26]. این داده‌ها در کنار زمان انتهابه‌انتهای برنامه، مبنای تحلیل کارایی قرار می‌گیرند.")

    add_heading(doc, "۲-۲-۱۰ کش سطح قالب و اعتبارسنجی تصمیم", 3)
    add_body(doc, "کش نگاشت یک کلید پایدار به تصمیم معتبر قبلی است. FlexLog برای کاهش فراخوانی LLM از کش و بازیابی استفاده می‌کند و AdaptiveLog با واگذاری نمونه‌های مطمئن به مدل کوچک‌تر، تعداد فراخوانی مدل بزرگ را تا ۷۳ درصد کاهش داده است [10، 12]. LogParser-LLM نیز با تجمیع الگوها در درخت پیشوند از تکرار فراخوانی LLM برای پیام‌های مشابه جلوگیری می‌کند [18]. این مطالعات نشان می‌دهند کنترل هزینه استنتاج بخشی مستقل از کیفیت طبقه‌بندی است.")
    add_body(doc, "در پژوهش حاضر کلید کش از نسخه پرامپت، نام و نسخه مدل، قالب نرمال‌شده و در صورت فعال بودن، مؤلفه، سطح شدت و دسته پیام ساخته می‌شود. بنابراین با تغییر پرامپت یا مدل، تصمیم قبلی به‌اشتباه بازاستفاده نمی‌شود. فقط پاسخ معتبر ۰ یا ۱ ذخیره می‌شود؛ INVALID، خطای ارتباط یا پاسخ ناقص در کش قرار نمی‌گیرد. نرخ برخورد کش، نرخ عدم برخورد و تعداد کلیدهای یکتا همراه با منبع تصمیم ثبت می‌شود تا اثر کش از اثر مدل و قواعد جدا شود.")

    add_heading(doc, "۲-۲-۱۱ دو معماری Hybrid و Prompt-only", 3)
    add_body(doc, "در معماری Hybrid این پژوهش، خط پس از استخراج فیلدها و نرمال‌سازی ابتدا در کش جست‌وجو می‌شود. در نبود تصمیم کش‌شده، مجموعه‌ای محدود از قواعد قطعی و با اطمینان بالا بررسی می‌گردد. اگر هیچ قاعده‌ای تصمیم ندهد، قالب به Qwen ارسال و پاسخ پس از اعتبارسنجی ذخیره می‌شود. منبع هر تصمیم به‌صورت cache، guard یا llm ثبت می‌شود. در نتیجه Hybrid یک آبشار تصمیم است و هدف آن سنجش اثر قاعده قطعی پیش از LLM است.")
    add_body(doc, "در معماری Prompt-only همان نرمال‌سازی، کلید کش، مدل و قرارداد خروجی حفظ می‌شود، اما guard قطعی غیرفعال است. دانش دامنه و قواعد تشخیص در متن پرامپت بیان می‌شوند و Qwen برای هر قالبِ فاقد کش تصمیم می‌گیرد. مقایسه منصفانه مستلزم ثابت ماندن مجموعه داده، ترتیب ورودی، نسخه مدل، گزینه‌های Ollama، نرمال‌سازی و سیاست اعتبارسنجی است. تفاوت این تعریف با Hybrid Prompting آن است که روش پارک و همکاران شواهد سطح رخداد، شباهت توالی و نمونه عادی مشابه را در یک پرامپت ترکیب می‌کند و طبقه‌بندی آن در سطح توالی است [9].")

    add_heading(doc, "۲-۲-۱۲ معیارهای ارزیابی", 3)
    add_body(doc, "برای مسئله دودویی، TP تعداد ناهنجاری‌های درست کشف‌شده، TN تعداد خطوط عادیِ درست، FP خطوط عادی با برچسب ناهنجار و FN ناهنجاری‌های ازدست‌رفته است. دقت کل نسبت همه تصمیم‌های درست به کل نمونه‌هاست؛ Precision سهم تشخیص‌های ناهنجار درست از همه پیش‌بینی‌های ناهنجار؛ Recall سهم ناهنجاری‌های کشف‌شده از همه ناهنجاری‌های واقعی؛ و F1 میانگین هارمونیک Precision و Recall است [3، 9].")
    add_body(doc, "Accuracy = (TP + TN) / (TP + TN + FP + FN)", center=True, ltr=True, size=12)
    add_body(doc, "Precision = TP / (TP + FP)     Recall = TP / (TP + FN)     F1 = 2 × Precision × Recall / (Precision + Recall)", center=True, ltr=True, size=12)
    add_body(doc, "به دلیل نامتوازن بودن BGL، F1، Precision و Recall باید در کنار Accuracy گزارش شوند. افزون بر این چهار معیار، نرخ INVALID، زمان پاسخ، تعداد فراخوانی LLM، منبع تصمیم و آمار برخورد کش برای مقایسه دو معماری ضروری است. گزارش این شاخص‌ها کمک می‌کند افزایش سرعت ناشی از کش یا guard با کیفیت طبقه‌بندی اشتباه نشود.")

    add_heading(doc, "۲-۳ پیشینه تحقیق", 2)
    add_heading(doc, "۲-۳-۱ روش‌های ترتیبی و پیش‌آموزش‌دیده", 3)
    add_body(doc, "Du و همکاران DeepLog را به‌عنوان یک مدل LSTM برای یادگیری توالی رخدادهای عادی، تشخیص انحراف و به‌روزرسانی برخط ارائه کردند [1]. Guo و همکاران در LogBERT از BERT و دو هدف خودنظارتی برای یادگیری الگوی توالی‌های عادی بهره گرفتند و روش را روی HDFS، BGL و Thunderbird ارزیابی کردند [2]. Chen و Liao در BERT-Log بازنمایی BERT را با طبقه‌بند تمام‌متصل ترکیب کردند و در پروتکل خود F1 برابر ۹۹٫۳ درصد برای HDFS و ۹۹٫۴ درصد برای BGL گزارش دادند [3].")
    add_body(doc, "LogFiT از تنظیم مدل زبانی پیش‌آموزش‌دیده با وظیفه توکن ماسک‌شده روی لاگ‌های عادی و قاعده top-k برای تشخیص استفاده می‌کند و روی HDFS، BGL و Thunderbird آزمایش شده است [13]. ADALog وابستگی به توالی و تجزیه را حذف و با تنظیم DistilBERT روی لاگ‌های عادی، احتمال بازسازی توکن و آستانه صدکی تطبیقی، تصمیم سطح خط تولید می‌کند [17]. این دو مطالعه نشان می‌دهند مدل زبانی پیش‌آموزش‌دیده می‌تواند بدون برچسب ناهنجاری به‌کار رود، اما همچنان به مرحله آموزش یا تنظیم روی داده مقصد نیاز دارد.")

    add_heading(doc, "۲-۳-۲ روش‌های مبتنی بر آموزش یا تنظیم LLM", 3)
    add_body(doc, "Han، Yuan و Trabelsi در LogGPT ابتدا GPT را برای پیش‌بینی رخداد بعدی آموزش و سپس با یادگیری تقویتی به هدف تشخیص ناهنجاری هم‌راستا کردند؛ ارزیابی روی HDFS، BGL و Thunderbird انجام شد [4]. Guan و همکاران در LogLLM از BERT برای بردارهای معنایی، یک نگاشت بین بازنمایی‌ها و LLaMA به‌عنوان طبقه‌بند استفاده کردند و یک فرایند آموزشی چندمرحله‌ای را روی چهار مجموعه HDFS، BGL، Liberty و Thunderbird گزارش کردند [5].")
    add_body(doc, "Zhang و همکاران در LLM-LADE سه مرحله افزایش داده، تنظیم پارامترکارا و بهینه‌سازی برخط را برای تشخیص و تبیین هم‌زمان ترکیب کردند؛ در پیاده‌سازی از GPT-4 برای تولید نمونه و Llama3-8B با LoRA برای مدل مقصد استفاده شد [15]. Lim، Zhu و Pang نیز LoRA و ReFT را روی RoBERTa، GPT-2 و Llama-3 از نظر پایداری، کارایی نمونه و تعمیم بین مجموعه‌ها بررسی کردند [14]. این دسته کیفیت مدل را از طریق تغییر پارامترها یا نگاشت بازنمایی افزایش می‌دهد، اما هزینه آموزش و وابستگی به داده تنظیم را وارد فرایند می‌کند.")

    add_heading(doc, "۲-۳-۳ روش‌های مبتنی بر پرامپت", 3)
    add_body(doc, "Qi و همکاران در مطالعه LogGPT امکان استفاده از ChatGPT برای تشخیص ناهنجاری BGL و Spirit را با پرامپت‌های مختلف بررسی کردند و علاوه بر برچسب، تبیین متنی را ارزیابی کردند [6]. Liu و همکاران در LogPrompt مجموعه‌ای از راهبردهای پرامپت برای تجزیه و تشخیص ارائه و آن را بدون آموزش درون‌دامنه روی ۹ مجموعه داده آزمایش کردند؛ نتایج مقاله بهبود تا ۵۵٫۹ درصد نسبت به روش‌های آموزش‌دیده و امتیاز تفسیرپذیری ۴٫۴۲ از ۵ در ارزیابی متخصصان را گزارش می‌کند [7]. فایل سال ۲۰۲۳ موجود در مجموعه منابع، نسخه اولیه همین خط پژوهش است و در جدول به‌عنوان مطالعه مستقل تکرار نشده است.")
    add_body(doc, "Park، Choi و Lee در Hybrid Prompting سه نوع شواهد را در پرامپت ترکیب کردند: مشاهده رخدادهای دیده‌نشده، شباهت یا توزیع توالی و نزدیک‌ترین نمونه عادی. روش با GPT-4o-mini و مدل تعبیه Qwen3-Embedding-8B روی HDFS، BGL و Thunderbird آزمایش شد و میانگین بهبود F1 برابر ۹٫۹۳ واحد درصد نسبت به خط‌پایه‌ها گزارش شد؛ F1 روش روی BGL در پروتکل توالی یک‌دقیقه‌ای ۷۷٫۸۱ درصد بود [9]. این نتیجه به دلیل سطح توالی و نمونه متوازن با طبقه‌بندی خط‌به‌خط این پژوهش قابل مقایسه عددی مستقیم نیست.")

    add_heading(doc, "۲-۳-۴ بازیابی، روش‌های ترکیبی و کاهش هزینه", 3)
    add_body(doc, "RAGLog پایگاه برداری از لاگ‌های عادی می‌سازد و با بازیابی نمونه‌های مرتبط، LLM را در تشخیص بدون آموزش اختصاصی هدایت می‌کند؛ ارزیابی آن روی BGL و Thunderbird انجام شده است [8]. AdaptiveLog همکاری مدل کوچک و بزرگ را به‌کار می‌گیرد: عدم‌قطعیت مدل کوچک زمان فراخوانی LLM را تعیین می‌کند و نمونه‌های دشوار مشابه به پرامپت افزوده می‌شوند؛ مقاله کاهش هزینه فراخوانی LLM تا ۷۳ درصد را گزارش می‌کند [12].")
    add_body(doc, "Hadadi و همکاران در FlexLog خروجی درخت تصمیم، KNN و شبکه پیش‌خور را با Mistral، کش و بازیابی ترکیب کردند و روش را روی مجموعه‌های ناپایدار ADFA، LOGEVOL و داده‌های مصنوعی HDFS ارزیابی کردند. مقاله دست‌کم ۱٫۲ واحد درصد بهبود F1 با ۶۲٫۸۷ واحد درصد داده برچسب‌خورده کمتر و در برخی تنظیم‌ها تا ۱۳ واحد درصد بهبود را گزارش می‌کند [10]. مطالعه دیگر همین گروه مقایسه تنظیم GPT-3 و مهندسی پرامپت GPT-4 را روی نسخه‌های ناپایدار LOGEVOL انجام داد و برتری تنظیم اختصاصی را در پروتکل خود گزارش کرد [11].")
    add_body(doc, "SemiRALD از ChatGPT و یادگیری درون‌متنی برای تجزیه، سپس RoBERTa و BiLSTM توجه‌محور برای تشخیص نیمه‌نظارتی استفاده می‌کند و روی HDFS و BGL میانگین بهبود F1 برابر ۷٫۳ و ۸٫۲ درصد را نسبت به دو گروه خط‌پایه گزارش کرده است [16]. بنابراین واژه «ترکیبی» در این گروه می‌تواند به ترکیب چند مدل، ترکیب بازیابی و LLM، یا ترکیب مرحله تجزیه و طبقه‌بندی اشاره کند و باید برای هر مطالعه به‌طور عملیاتی تعریف شود.")

    add_heading(doc, "۲-۳-۵ تجزیه، انتقال، تبیین و کاربرد عملی", 3)
    add_body(doc, "LogParser-LLM و LogPPT مستقیماً آشکارساز ناهنجاری نیستند، اما مرحله ورودی تحلیل لاگ را هدف می‌گیرند. LogPPT با ۳۲ نمونه برچسب‌خورده روی ۱۶ مجموعه به میانگین بیش از ۰٫۹ برای Group Accuracy و Parsing Accuracy رسیده و LogParser-LLM با درخت پیشوند، فراخوانی LLM را در مقیاس میلیون‌ها پیام محدود کرده است [18، 19]. این نتایج برای نرمال‌سازی و بازاستفاده تصمیم سطح قالب اهمیت دارند.")
    add_body(doc, "LogSynergy با تفسیر رخداد به‌وسیله LLM و استخراج ویژگی یکپارچه بین سامانه‌ها، انتقال آشکارساز به سامانه جدید را بررسی کرده و با ۵۰۰۰ توالی برچسب‌خورده F1 بیش از ۸۹ درصد روی داده واقعی و بیش از ۸۳ درصد روی داده عمومی گزارش کرده است [20]. LogLAA تجزیه تطبیقی، ویژگی‌های شمارشی، ترتیبی و معنایی، CNN-LSTM دو‌توجهی و LLM برای تولید توضیح را در یک چارچوب یکپارچه قرار داده است [21].")
    add_body(doc, "در مطالعه کاربرد واقعی AIOps، De la Cruz Cabello و همکاران LogBERT را فقط با توالی‌های عادی Linux syslog آموزش و به‌صورت محلی مستقر کردند. پنجره ۱۵ ثانیه‌ای با هم‌پوشانی ۱۰ ثانیه بهترین موازنه گزارش‌شده بین اثربخشی و تأخیر بود و ارزیابی کیفی متخصصان کاربرد عملی روش را تأیید کرد [27]. این مطالعه بر محرمانگی و استقرار محلی تأکید دارد، ولی طبقه‌بند آن LogBERT آموزش‌دیده است، نه یک LLM دستورمحور بدون تنظیم.")

    add_heading(doc, "۲-۴ دسته‌بندی، جدول مقایسه و تحلیل", 2)
    add_body(doc, "براساس مرور نظام‌مند حوزه لاگ و رده‌بندی نقش LLM در تشخیص ناهنجاری، مطالعات بررسی‌شده را می‌توان در چهار مسیر اصلی قرار داد: مدل‌سازی توالی، مدل زبانی پیش‌آموزش‌دیده با آموزش مقصد، LLM با تنظیم یا انتقال، و LLM مبتنی بر پرامپت، بازیابی یا همکاری مدل‌ها [22، 24]. شکل ۲-۱ رابطه این مسیرها و مؤلفه‌هایی را که در پژوهش حاضر کنار هم قرار می‌گیرند نشان می‌دهد.")
    add_relationship_figure(doc)
    add_caption(doc, "شکل ۲-۱: دسته‌بندی پژوهش‌های تشخیص ناهنجاری لاگ و جایگاه پژوهش حاضر", "F")
    add_body(doc, "در شکل ۲-۱، مسیر بالایی تحول روش‌ها را نمایش می‌دهد و بخش پایینی مشخصات عملیاتی پژوهش حاضر را از ادبیات متمایز می‌کند. قرارگرفتن یک مطالعه در یک دسته به معنای نبود اجزای دسته دیگر نیست؛ برای نمونه FlexLog هم مدل‌های یادگیری ماشین و هم LLM، بازیابی و کش را به‌کار می‌گیرد [10] و SemiRALD تجزیه با ChatGPT را به مدل ترتیبی پیوند می‌دهد [16].")

    add_caption(doc, "جدول ۲-۱: مقایسه مطالعات منتخب در تشخیص ناهنجاری و تحلیل لاگ", "T")
    headers1 = ["مطالعه", "سال", "روش/مدل", "داده‌ها", "واحد تحلیل", "نوع انطباق", "یافته یا کارکرد گزارش‌شده"]
    rows1 = [
        ("DeepLog [1]", "۲۰۱۷", "LSTM و پیش‌بینی رخداد بعدی", "HDFS، OpenStack", "توالی", "آموزش روی عادی", "تشخیص برخط و به‌روزرسانی افزایشی"),
        ("LogBERT [2]", "۲۰۲۱", "BERT با دو هدف خودنظارتی", "HDFS، BGL، Thunderbird", "توالی", "آموزش روی عادی", "یادگیری زمینه دوطرفه و الگوی مشترک عادی"),
        ("BERT-Log [3]", "۲۰۲۲", "BERT و طبقه‌بند متصل", "HDFS، BGL", "توالی/پنجره", "تنظیم نظارت‌شده", "F1 گزارش‌شده ۹۹٫۴٪ روی BGL"),
        ("LogGPT (Han) [4]", "۲۰۲۳", "GPT، پیش‌بینی رخداد و RL", "HDFS، BGL، Thunderbird", "توالی", "آموزش و RL", "هم‌راستاسازی هدف تولید با تشخیص"),
        ("LogLLM [5]", "۲۰۲۴", "BERT + نگاشت + LLaMA", "چهار مجموعه شامل BGL", "توالی", "آموزش چندمرحله‌ای", "ترکیب بازنمایی BERT با طبقه‌بند LLaMA"),
        ("LogGPT (Qi) [6]", "۲۰۲۳", "ChatGPT و پرامپت", "BGL، Spirit", "خط/پنجره", "بدون تنظیم", "بررسی پرامپت، دانش دامنه و تبیین"),
        ("LogPrompt [7]", "۲۰۲۴", "راهبردهای پرامپت و تبیین", "۹ مجموعه لاگ", "قالب/توالی", "بدون آموزش درون‌دامنه", "تا ۵۵٫۹٪ بهبود و امتیاز تبیین ۴٫۴۲/۵"),
        ("RAGLog [8]", "۲۰۲۳", "RAG و پایگاه برداری عادی", "BGL، Thunderbird", "توالی", "بدون تنظیم", "بازیابی نمونه عادی برای تصمیم معنایی"),
        ("Hybrid Prompting [9]", "۲۰۲۵", "پرامپت رخداد + توالی + نمونه", "HDFS، BGL، Thunderbird", "توالی یک‌دقیقه‌ای", "بدون تنظیم طبقه‌بند", "میانگین ۹٫۹۳ واحد درصد بهبود F1"),
        ("FlexLog [10]", "۲۰۲۴/۲۵", "ML ensemble + Mistral + RAG + cache", "چهار مجموعه ناپایدار", "توالی", "داده کم/درون‌متنی", "کاهش داده برچسب‌خورده و استنتاج کمتر از یک ثانیه"),
        ("Unstable Logs GPT [11]", "۲۰۲۴", "GPT-3 تنظیم‌شده و GPT-4 پرامپت", "LOGEVOL-Hadoop", "توالی", "تنظیم/پرامپت", "برتری تنظیم در پروتکل لاگ ناپایدار"),
        ("AdaptiveLog [12]", "۲۰۲۵", "مدل کوچک + LLM + بازیابی", "چند وظیفه لاگ", "وابسته به وظیفه", "همکاری تطبیقی", "تا ۷۳٪ کاهش هزینه فراخوانی LLM"),
        ("LogFiT [13]", "۲۰۲۵", "مدل ماسک‌شده و top-k", "HDFS، BGL، Thunderbird", "توالی", "تنظیم روی عادی", "تشخیص خودنظارتی بدون ناهنجاری برچسب‌خورده"),
        ("PEFT-LAD [14]", "۲۰۲۵", "LoRA/ReFT روی چند LM", "HDFS، BGL، Spirit، Thunderbird", "توالی", "تنظیم پارامترکارا", "تحلیل پایداری، کارایی نمونه و انتقال"),
        ("LLM-LADE [15]", "۲۰۲۵", "افزایش داده + LoRA + بهینه‌سازی برخط", "مجموعه‌های عمومی لاگ", "پیام/توالی", "PEFT", "تشخیص و تبیین هم‌زمان"),
        ("SemiRALD [16]", "۲۰۲۵", "ChatGPT parsing + RoBERTa + BiLSTM", "HDFS، BGL", "توالی", "نیمه‌نظارتی", "بهبود متوسط F1 به‌ویژه با داده کم"),
        ("ADALog [17]", "۲۰۲۵", "DistilBERT MLM و آستانه تطبیقی", "BGL، Thunderbird، Spirit", "خط", "تنظیم روی عادی", "بدون parser و مستقل از توالی"),
        ("LogParser-LLM [18]", "۲۰۲۴", "LLM + درخت پیشوند", "Loghub-2k، LogPub", "خط/قالب", "بدون آموزش برچسبی", "کاهش فراخوانی؛ F1 گروه‌بندی ۹۰٫۶٪"),
        ("LogPPT [19]", "۲۰۲۳", "پرامپت‌تیونینگ RoBERTa", "۱۶ مجموعه", "خط/قالب", "Few-shot؛ ۳۲ نمونه", "بیش از ۰٫۹ میانگین دقت گروه/تجزیه"),
        ("LogSynergy [20]", "۲۰۲۵", "تفسیر LLM و انتقال", "عمومی و واقعی", "توالی", "انتقال با ۵۰۰۰ برچسب", "F1 بیش از ۸۳٪ عمومی و ۸۹٪ واقعی"),
        ("LogLAA [21]", "۲۰۲۶", "parser + CNN-LSTM + LLM explanation", "چند مجموعه عمومی", "توالی", "آموزش چندجزئی", "تشخیص یکپارچه و تولید توضیح"),
        ("AIOps واقعی [27]", "۲۰۲۶", "LogBERT محلی", "Linux syslog واقعی", "پنجره زمانی", "خودنظارتی", "استقرار محلی تحت محدودیت محرمانگی"),
    ]
    add_data_table(doc, headers1, rows1, [Cm(2.5), Cm(1.0), Cm(2.8), Cm(2.3), Cm(1.7), Cm(2.1), Cm(3.7)], font_size=8.1)
    add_body(doc, "جدول ۲-۱ نشان می‌دهد که BGL در بسیاری از مطالعات حضور دارد، اما پروتکل‌های آن یکسان نیستند: برخی پیام را مستقیماً طبقه‌بندی می‌کنند، برخی پنجره زمانی یا توالی مبتنی بر گره می‌سازند و برخی فقط بخشی نمونه‌گیری‌شده را ارزیابی می‌کنند [2، 3، 6، 9، 17]. بنابراین استفاده مشترک از نام BGL به‌تنهایی هم‌ارزی آزمایش‌ها را تضمین نمی‌کند.")
    add_body(doc, "از نظر انطباق مدل، چهار حالت قابل مشاهده است: آموزش از ابتدا یا روی عادی، تنظیم نظارت‌شده یا پارامترکارا، پرامپت بدون تغییر وزن و همکاری چند مدل. پژوهش حاضر در حالت پرامپت بدون تغییر وزن قرار می‌گیرد، ولی با نرمال‌سازی و کش قطعی در سطح قالب و دو مسیر تصمیم کنترل‌شده تکمیل می‌شود. این ترکیب در جدول ۲-۲ با نزدیک‌ترین مطالعات مقایسه شده است.")

    table2_caption = add_caption(doc, "جدول ۲-۲: مقایسه سازوکار مطالعات نزدیک با مؤلفه‌های پژوهش حاضر", "T")
    table2_caption.paragraph_format.page_break_before = True
    headers2 = ["مطالعه", "سازوکار اصلی", "سطح تصمیم", "کنترل هزینه", "نوع استقرار", "تفاوت با پژوهش حاضر"]
    rows2 = [
        ("LogGPT (Qi) [6]", "پرامپت ChatGPT", "خط/پنجره", "نمونه‌گیری به علت محدودیت API", "خدمت تجاری", "فاقد Qwen محلی، کش قالب و مقایسه guard"),
        ("LogPrompt [7]", "راهبرد پرامپت و تبیین", "قالب/توالی", "بدون کش عملیاتیِ محور مقایسه", "LLMهای مختلف", "تمرکز بر راهبرد پرامپت، نه آبشار guard و cache"),
        ("RAGLog [8]", "بازیابی لاگ عادی", "توالی", "کاهش زمینه با retrieval", "LLM بیرونی", "پایگاه برداری نمونه است، نه کش تصمیم قالب"),
        ("Hybrid Prompting [9]", "ترکیب شواهد در پرامپت", "توالی", "بازیابی نزدیک‌ترین نمونه", "GPT-4o-mini", "Hybrid به معنای ترکیب پرامپت و در سطح توالی است"),
        ("FlexLog [10]", "ensemble + LLM + RAG + cache", "توالی", "کش و واگذاری به ML", "Mistral", "داده‌های ناپایدار و ترکیب چند طبقه‌بند؛ نه BGL خط‌محور"),
        ("AdaptiveLog [12]", "مدل کوچک + LLM", "چند وظیفه", "فراخوانی براساس عدم‌قطعیت", "همکاری دو مدل", "به مدل کوچک آموزش‌دیده نیاز دارد"),
        ("ADALog [17]", "MLM و آستانه تطبیقی", "خط", "LLM مولد فراخوانی نمی‌شود", "مدل محلی", "نیازمند fine-tuning DistilBERT؛ فاقد پرامپت دودویی"),
        ("LogParser-LLM [18]", "درخت پیشوند + LLM", "قالب", "تجمیع در درخت و فراخوانی محدود", "LLM قابل انتخاب", "هدف تجزیه است؛ تصمیم ناهنجاری تولید نمی‌کند"),
        ("AIOps واقعی [27]", "LogBERT محلی", "پنجره", "تنظیم پنجره و مدل خودنظارتی", "کاملاً محلی", "مدل آموزش‌دیده روی syslog، نه Qwen دستورمحور"),
        ("پژوهش حاضر", "دو روش Hybrid و Prompt-only", "خط BGL", "نرمال‌سازی + کش تصمیم قالب", "Qwen2.5:7B روی Ollama", "مقایسه کنترل‌شده guard بیرونی با قاعده درون پرامپت"),
    ]
    add_data_table(doc, headers2, rows2, [Cm(2.5), Cm(3.1), Cm(1.8), Cm(3.1), Cm(2.5), Cm(3.3)], font_size=8.5)
    add_body(doc, "مطالعات نزدیک هرکدام بخشی از مسئله را پوشش داده‌اند: LogPrompt و LogGPT بر رفتار پرامپت، RAGLog بر بازیابی، FlexLog بر ترکیب کش و مدل‌های کلاسیک، AdaptiveLog بر واگذاری تطبیقی و LogParser-LLM بر کاهش فراخوانی در سطح قالب تمرکز دارند [6–10، 12، 18]. در مجموعه منابع بررسی‌شده، مقایسه کنترل‌شده دو مسیر با مدل، داده، نرمال‌سازی و کش یکسان—یکی با guard قطعی بیرون مدل و دیگری با انتقال همان دانش به پرامپت—به‌صورت مستقیم گزارش نشده است.")

    add_heading(doc, "۲-۵ شکاف تحقیقاتی و جایگاه پژوهش حاضر", 2)
    add_body(doc, "نخستین شکاف به ناهمگونی واحد تحلیل و پروتکل BGL مربوط است. نتایج خط‌محور، پنجره مبتنی بر گره، پنجره یک‌دقیقه‌ای و توالی پنج‌دقیقه‌ای در ادبیات با یک نام مجموعه داده گزارش شده‌اند [2، 3، 9، 17]. این وضعیت نیازمند تعریف صریح واحد تصمیم، جلوگیری از ورود برچسب BGL به ورودی و گزارش ماتریس اغتشاش در سطح همان واحد است. پژوهش حاضر واحد تصمیم را خط تعیین می‌کند و برچسب را فقط برای ارزیابی نگه می‌دارد.")
    add_body(doc, "دومین شکاف، تفکیک اثر کیفیت تشخیص از هزینه استنتاج است. برخی مطالعات کش، بازیابی یا مدل کوچک را برای کاهش تماس با LLM به‌کار برده‌اند [10، 12، 18]، اما نرخ برخورد کش، منبع هر تصمیم، زمان پاسخ و نرخ INVALID همیشه در کنار Accuracy، Precision، Recall و F1 گزارش نمی‌شود. در پژوهش حاضر این متغیرها بخشی از خروجی آزمایش‌اند تا معلوم شود هر بهبود از مدل، guard یا بازاستفاده تصمیم ناشی شده است.")
    add_body(doc, "سومین شکاف به جایگاه قواعد دامنه مربوط است. در پژوهش‌های پرامپتی، دانش در متن دستور یا نمونه‌ها قرار می‌گیرد [6، 7، 9]؛ در روش‌های قاعده‌محور تصمیم خارج مدل است؛ و در سامانه‌های ترکیبی، مدل کلاسیک یا بازیابی نقش دروازه را دارد [10، 12]. مقایسه مستقیم قاعده قطعی خارج از LLM با همان منطق به‌صورت دستور درون پرامپت، تحت شرایط ثابت، در مجموعه مطالعات مرورشده مشاهده نشد. دو روش Hybrid و Prompt-only برای سنجش همین تمایز تعریف شده‌اند.")
    add_body(doc, "چهارمین شکاف، محدود بودن شواهد درباره یک LLM متن‌باز متوسط در استقرار کاملاً محلی و بدون تنظیم دقیق برای طبقه‌بندی خط‌به‌خط BGL است. بخشی از مطالعات از خدمات تجاری مانند ChatGPT یا GPT-4o-mini استفاده کرده‌اند [6، 9]، بخشی LLaMA یا Mistral را آموزش یا تنظیم کرده‌اند [5، 10، 15] و مطالعه محلی سال ۲۰۲۶ از LogBERT خودنظارتی بهره گرفته است [27]. پژوهش حاضر Qwen2.5:7B-Instruct را با Ollama محلی و وزن‌های ثابت ارزیابی می‌کند [25، 26].")
    add_body(doc, "برآیند این شکاف‌ها جایگاه پژوهش را چنین مشخص می‌کند: یک خط لوله بازتولیدپذیر برای BGL شامل استخراج و حذف برچسب، نرمال‌سازی حساس به معنا، کش نسخه‌بندی‌شده در سطح قالب، خروجی JSON اعتبارسنجی‌شده و ثبت منبع تصمیم؛ سپس مقایسه دو معماری که فقط در محل اعمال دانش قطعی تفاوت دارند. این صورت‌بندی ادعای برتری پیشینی هیچ مسیر را مطرح نمی‌کند و نتیجه باید از آزمایش کامل و یکسان دو روش به‌دست آید.")

    add_heading(doc, "۲-۶ جمع‌بندی", 2)
    add_body(doc, "در این فصل مفاهیم لاگ، ناهنجاری، واحد تحلیل، BGL، تجزیه و نرمال‌سازی، مدل‌های ترتیبی، مدل‌های زبانی پیش‌آموزش‌دیده، LLM، Qwen2.5، مهندسی پرامپت، استنتاج محلی، کش و معیارهای ارزیابی تعریف شد. مرور پیشینه نشان داد مسیر پژوهش از پیش‌بینی توالی و مدل‌های BERT به تنظیم LLM، پرامپت، بازیابی و سامانه‌های چندجزئی گسترش یافته است [1–24].")
    add_body(doc, "مقایسه مطالعات نشان داد که نوع واحد تحلیل، نیاز یا عدم نیاز به آموزش، سازوکار کنترل هزینه و محل استقرار برای تفسیر نتایج ضروری‌اند. شکاف استخراج‌شده از منابع، نبود مقایسه مستقیم و کنترل‌شده بین guard قطعی بیرون LLM و انتقال دانش آن به پرامپت، همراه با کش قالب و Qwen محلی در طبقه‌بندی خط‌به‌خط BGL است. فصل بعد باید معماری نرم‌افزار، قواعد نرمال‌سازی، ساختار پرامپت‌ها، کلید کش، تنظیمات Ollama، پروتکل اجرای دو روش و شیوه محاسبه معیارها را به‌صورت عملیاتی تشریح کند.")

    add_heading(doc, "۲-۷ منابع فصل دوم", 2)
    references = [
        "[1] Du, M.; Li, F.; Zheng, G.; Srikumar, V. DeepLog: Anomaly Detection and Diagnosis from System Logs through Deep Learning. ACM CCS, 2017. DOI: 10.1145/3133956.3134015.",
        "[2] Guo, H.; Yuan, S.; Wu, X. LogBERT: Log Anomaly Detection via BERT. arXiv:2103.04475, 2021.",
        "[3] Chen, S.; Liao, H. BERT-Log: Anomaly Detection for System Logs Based on Pre-trained Language Model. Applied Artificial Intelligence, 36(1), 2022. DOI: 10.1080/08839514.2022.2145642.",
        "[4] Han, X.; Yuan, S.; Trabelsi, M. LogGPT: Log Anomaly Detection via GPT. 2023.",
        "[5] Guan, W.; Cao, J.; Qian, S.; Gao, J.; Ouyang, C. LogLLM: Log-based Anomaly Detection Using Large Language Models. 2024.",
        "[6] Qi, J.; Huang, S.; Luan, Z.; Fung, C.; Yang, H.; Qian, D. LogGPT: Exploring ChatGPT for Log-Based Anomaly Detection. 2023.",
        "[7] Liu, Y.; Tao, S.; Meng, W.; Wang, J.; Ma, W.; Chen, Y.; Zhao, Y.; Yang, H.; Jiang, Y. Prompt Engineering Towards Zero-Shot and Interpretable Log Analysis. 2024. Earlier version: Interpretable Online Log Analysis Using Large Language Models with Prompt Strategies, 2023.",
        "[8] Pan, J.; Wong, S. L.; Yuan, Y. RAGLog: Log Anomaly Detection using Retrieval Augmented Generation. 2023.",
        "[9] Park, J.; Choi, E.; Lee, J. Hybrid Prompting for LLM-based Log Anomaly Detection. Journal of KIISE, 52(12), 2025. DOI: 10.5626/JOK.2025.52.12.1067.",
        "[10] Hadadi, F.; Xu, Q.; Bianculli, D.; Briand, L. LLM meets ML: Data-efficient Anomaly Detection on Unstable Logs. 2024–2025.",
        "[11] Hadadi, F.; Xu, Q.; Bianculli, D.; Briand, L. Anomaly Detection on Unstable Logs with GPT Models. 2024.",
        "[12] Ma, L.; Yang, W.; Li, Y.; Fei, B.; Zhou, M.; Li, S.; Jiang, S.; Xu, B.; Xiao, Y. AdaptiveLog: An Adaptive Log Analysis Framework with the Collaboration of Large and Small Language Model. arXiv:2501.11031, 2025.",
        "[13] Almodovar, C.; Sabrina, F.; Karimi, S.; Azad, S. LogFiT: Log Anomaly Detection using Fine-Tuned Language Models. 2025.",
        "[14] Lim, Y. F.; Zhu, J.; Pang, G. Adapting Large Language Models for Parameter-Efficient Log Anomaly Detection. 2025.",
        "[15] Zhang, Z.; Li, S.; Zhang, L.; Ye, J.; Hu, C.; Yan, L. LLM-LADE: Large language model-based log anomaly detection with explanation. Knowledge-Based Systems, 326, Article 114064, 2025.",
        "[16] Sun, Y.; Keung, J. W.; Yang, Z.; Liu, S.; Yu, H. K. SemiRALD: A semi-supervised hybrid language model for robust Anomalous Log Detection. Information and Software Technology, 183, 107743, 2025. DOI: 10.1016/j.infsof.2025.107743.",
        "[17] Pospieszny, P.; Mormul, W.; Szyndler, K.; Kumar, S. ADALog: Adaptive Unsupervised Anomaly Detection in Logs with Self-attention Masked Language Model. arXiv:2505.13496, 2025.",
        "[18] Zhong, A.; Mo, D.; Liu, G.; Liu, J.; Lu, Q.; Zhou, Q.; Wu, J.; Li, Q.; Wen, Q. LogParser-LLM: Advancing Efficient Log Parsing with Large Language Models. arXiv:2408.13727, 2024.",
        "[19] Le, V.-H.; Zhang, H. Log Parsing with Prompt-based Few-shot Learning. arXiv:2302.07435, 2023.",
        "[20] Sui, Y.; Wang, X.; Cui, T.; Xiao, T.; He, C.; Zhang, S.; Zhang, Y.; Yang, X.; Sun, Y.; Pei, D. Bridging the Gap: LLM-Powered Transfer Learning for Log Anomaly Detection in New Software Systems. 2025.",
        "[21] Gao, Y.; Luo, T.; Huang, K.; Tang, J.; Li, X. LogLAA: an adaptive integrated log anomaly analysis framework. Cybersecurity, 9:141, 2026. DOI: 10.1186/s42400-026-00573-8.",
        "[22] De la Cruz Cabello, M.; Prince Sales, T.; Machado, M. R. AIOps for log anomaly detection in the era of LLMs: A systematic literature review. Intelligent Systems with Applications, 28, Article 200608, 2025.",
        "[23] Liu, Y.; Pei, C.; Xu, L.; et al. OpsEval: A Comprehensive Benchmark Suite for Evaluating Large Language Models’ Capability in IT Operations Domain. FSE Companion, 2025. DOI: 10.1145/3696630.3728572.",
        "[24] Xu, R.; Ding, K. Large Language Models for Anomaly and Out-of-Distribution Detection: A Survey. arXiv:2409.01980, 2025.",
        "[25] Qwen Team. Qwen2.5 Technical Report. arXiv:2412.15115, 2024.",
        "[26] Ollama. API Introduction; Structured Outputs; Usage Metrics. Official documentation. Accessed 14 August 2026.",
        "[27] De la Cruz Cabello, M.; Prince Sales, T.; Machado, M. R. Log anomaly detection in AIOps: A real-world implementation using Large Language Models. Systems and Soft Computing, 8, Article 200475, 2026.",
    ]
    for ref in references:
        p = add_body(doc, ref, ltr=True, size=11)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_together = True


def normalize_new_content(doc, chapter):
    started = False
    for p in doc.paragraphs:
        if p._p is chapter._p:
            started = True
        if not started:
            continue
        if p.style.name in ("Heading 1", "Heading 2", "Heading 3", "Caption", "Caption1"):
            set_bidi(p, True)
            continue
        for r in p.runs:
            if r.text:
                current_size = r.font.size.pt if r.font.size else 14
                set_run_fonts(r, size=current_size)


def main():
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)
    else:
        shutil.copy2(BACKUP, DOCX)
    doc = Document(DOCX)
    chapter = remove_after_chapter_heading(doc)
    replace_front_fields(doc)
    add_chapter_content(doc)
    normalize_new_content(doc, chapter)
    doc.core_properties.comments = "نسخه کاری؛ فصل دوم بر پایه منابع مرورشده بازنویسی شد؛ فایل اصلی بدون تغییر است."
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
