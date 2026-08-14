from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.shared import RGBColor


ROOT = Path(r"D:\Programming\Thesis\LLMLogAnalyzer")
SOURCE = ROOT / "documents" / "پایان نامه ی مسعود دباغی.docx"
OUTPUT = ROOT / "documents" / "پایان نامه ی مسعود دباغی - نسخه کاری.docx"


def set_run_fonts(run, persian="B Nazanin", latin="Times New Roman", size=None, bold=None):
    run.font.name = latin
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._r.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), persian)
    fonts.set(qn("w:cs"), persian)


def set_style_fonts(style, persian="B Nazanin", latin="Times New Roman", size=14, bold=None):
    style.font.name = latin
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:eastAsia"), persian)
    fonts.set(qn("w:cs"), persian)


def set_bidi(paragraph, enabled=True):
    ppr = paragraph._p.get_or_add_pPr()
    bidi = ppr.find(qn("w:bidi"))
    if enabled and bidi is None:
        ppr.append(OxmlElement("w:bidi"))
    elif not enabled and bidi is not None:
        ppr.remove(bidi)


def set_keep(paragraph, keep_next=False, keep_lines=True, page_break_before=False):
    ppr = paragraph._p.get_or_add_pPr()
    for tag, enabled in (
        ("w:keepNext", keep_next),
        ("w:keepLines", keep_lines),
        ("w:pageBreakBefore", page_break_before),
    ):
        existing = ppr.find(qn(tag))
        if enabled and existing is None:
            ppr.append(OxmlElement(tag))
        elif not enabled and existing is not None:
            ppr.remove(existing)


def ensure_style(doc, name, base="Normal", style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, style_type)
        if base:
            style.base_style = doc.styles[base]
        return style


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_fonts(normal, size=14)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    ppr = normal.element.get_or_add_pPr()
    if ppr.find(qn("w:bidi")) is None:
        ppr.append(OxmlElement("w:bidi"))

    h1 = doc.styles["Heading 1"]
    set_style_fonts(h1, size=18, bold=True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(72)
    h1.paragraph_format.space_after = Pt(0)
    h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    h1.paragraph_format.keep_with_next = False
    h1.paragraph_format.page_break_before = True

    h2 = doc.styles["Heading 2"]
    set_style_fonts(h2, size=14, bold=True)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = doc.styles["Heading 3"]
    set_style_fonts(h3, size=14, bold=True)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    front = ensure_style(doc, "عنوان مقدماتی")
    set_style_fonts(front, size=16, bold=True)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.space_before = Pt(0)
    front.paragraph_format.space_after = Pt(12)

    label = ensure_style(doc, "برچسب")
    set_style_fonts(label, size=14, bold=True)
    label.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label.paragraph_format.space_before = Pt(6)
    label.paragraph_format.space_after = Pt(2)

    caption = ensure_style(doc, "Caption")
    set_style_fonts(caption, size=11, bold=False)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    for name, size, indent in (("TOC 1", 13, 0), ("TOC 2", 12, 0.5), ("TOC 3", 11, 1.0)):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        set_style_fonts(style, size=size)
        style.paragraph_format.right_indent = Cm(indent)
        style.paragraph_format.left_indent = Cm(0)
        style.paragraph_format.space_after = Pt(2)

    for name in ("footnote text", "FootNote"):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        set_style_fonts(style, size=12)


def set_text(paragraph, text, style=None, alignment=None, bold=None, size=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_fonts(run, size=size, bold=bold)
    if style:
        paragraph.style = style
    if alignment is not None:
        paragraph.alignment = alignment
    set_bidi(paragraph, True)
    return paragraph


def insert_after(paragraph, text, style):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    result = Paragraph(new_p, paragraph._parent)
    set_text(result, text, style=style)
    return result


def add_complex_field(paragraph, instruction, placeholder=""):
    paragraph.clear()
    begin = paragraph.add_run()
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin._r.append(begin_char)

    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    instr._r.append(instr_text)

    separate = paragraph.add_run()
    separate_char = OxmlElement("w:fldChar")
    separate_char.set(qn("w:fldCharType"), "separate")
    separate._r.append(separate_char)

    if placeholder:
        result = paragraph.add_run(placeholder)
        set_run_fonts(result, size=12)

    end = paragraph.add_run()
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end._r.append(end_char)
    set_bidi(paragraph, True)


def add_hidden_tc_field(paragraph, entry, identifier="T"):
    begin = paragraph.add_run()
    begin_char = OxmlElement("w:fldChar")
    begin_char.set(qn("w:fldCharType"), "begin")
    begin._r.append(begin_char)
    instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f'TC "{entry}" \\f {identifier} \\l "1"'
    instr._r.append(instr_text)
    end = paragraph.add_run()
    end_char = OxmlElement("w:fldChar")
    end_char.set(qn("w:fldCharType"), "end")
    end._r.append(end_char)


def next_paragraph(doc, paragraph):
    paragraphs = doc.paragraphs
    idx = next(i for i, item in enumerate(paragraphs) if item._p is paragraph._p)
    return paragraphs[idx + 1]


def format_front_matter(doc):
    front_titles = {
        "فهرست مطالب",
        "فهرست جدول ها",
        "فهرست نمودار ها",
        "فهرست شکل ها",
        "چکیده",
    }
    labels = {"کلمات کلیدی(فارسی)", "کلمات کلیدی(انگلیسی)"}
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in front_titles:
            paragraph.style = "عنوان مقدماتی"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_bidi(paragraph, True)
        elif text in labels:
            paragraph.style = "برچسب"
            set_bidi(paragraph, True)

    page_starts = (
        "معاونت پژوهش و فنآوري",
        "دانشگاه آزاد اسلامی",
        "بر خود لازم می‌دانم",
        "تقدیم به",
        "فهرست مطالب",
        "فهرست جدول ها",
        "فهرست نمودار ها",
        "فهرست شکل ها",
    )
    seen_university = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        should_break = False
        if text == "دانشگاه آزاد اسلامی":
            seen_university += 1
            should_break = True
        elif any(text.startswith(prefix) for prefix in page_starts if prefix != "دانشگاه آزاد اسلامی"):
            should_break = True
        if should_break:
            set_keep(paragraph, page_break_before=True)

    toc_title = next(p for p in doc.paragraphs if p.text.strip() == "فهرست مطالب")
    toc_body = next_paragraph(doc, toc_title)
    add_complex_field(toc_body, 'TOC \\o "1-3" \\h \\z \\u', "فهرست مطالب پس از بازشدن در Word به‌روزرسانی می‌شود.")

    lot_title = next(p for p in doc.paragraphs if p.text.strip() == "فهرست جدول ها")
    lot_body = next_paragraph(doc, lot_title)
    add_complex_field(lot_body, 'TOC \\f T \\h \\z', "فهرست جدول‌ها پس از بازشدن در Word به‌روزرسانی می‌شود.")

    abstract_title = next(p for p in doc.paragraphs if p.text.strip() == "چکیده")
    trim_blank_paragraphs_before(abstract_title)


def format_headings(doc):
    p = next(p for p in doc.paragraphs if p.text.strip().startswith("فصل اول: کلیات تحقیق"))
    remove_preceding_blank_page_breaks(p)
    set_text(p, "فصل اول: کلیات تحقیق", style="Heading 1", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
    set_keep(p, page_break_before=True)
    p.add_run().add_break(WD_BREAK.PAGE)
    intro = insert_after(p, "۱-۱ مقدمه", "Heading 2")
    set_keep(intro, keep_next=True)

    mapping = {
        "1-2 بيان مسأله": "۱-۲ بیان مسئله",
        "1-3 اهمیت و ضرورت انجام پژوهش": "۱-۳ اهمیت و ضرورت انجام پژوهش",
        "1-4 اهداف مشخص پژوهش": "۱-۴ اهداف مشخص پژوهش",
        "1-5 سؤال‌ها یا فرضیه‌های پژوهش": "۱-۵ سؤال‌ها و فرضیه‌های پژوهش",
        "سؤالات پژوهش": "سؤال‌های پژوهش",
        "فرضیه‌های پژوهش": "فرضیه‌های پژوهش",
        "فصل دوم: مروري بر ادبیات تحقیق و پیشینه تحقیق": "فصل دوم: مروری بر ادبیات و پیشینه تحقیق",
        "۲-۱ مقدمه": "۲-۱ مقدمه",
        "2-2 مرور ادبیات و سوابق مربوط": "۲-۲ مرور ادبیات و سوابق مرتبط",
    }
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        for prefix, replacement in mapping.items():
            if text.startswith(prefix):
                if replacement.startswith("فصل دوم"):
                    remove_preceding_blank_page_breaks(paragraph)
                    set_text(paragraph, replacement, style="Heading 1", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18)
                    set_keep(paragraph, page_break_before=True)
                    paragraph.add_run().add_break(WD_BREAK.PAGE)
                elif replacement in ("سؤال‌های پژوهش", "فرضیه‌های پژوهش"):
                    set_text(paragraph, replacement, style="Heading 3", bold=True, size=14)
                    set_keep(paragraph, keep_next=True)
                else:
                    set_text(paragraph, replacement, style="Heading 2", bold=True, size=14)
                    set_keep(paragraph, keep_next=True)
                break


def remove_preceding_blank_page_breaks(paragraph):
    trim_blank_paragraphs_before(paragraph)


def trim_blank_paragraphs_before(paragraph):
    prev = paragraph._p.getprevious()
    while prev is not None and prev.tag == qn("w:p"):
        texts = prev.findall(".//" + qn("w:t"))
        if any((node.text or "").strip() for node in texts):
            break
        before = prev.getprevious()
        for br in list(prev.findall(".//" + qn("w:br"))):
            if br.get(qn("w:type")) == "page":
                br.getparent().remove(br)
        ppr = prev.find(qn("w:pPr"))
        has_section = ppr is not None and ppr.find(qn("w:sectPr")) is not None
        if ppr is not None:
            marker = ppr.find(qn("w:pageBreakBefore"))
            if marker is not None:
                ppr.remove(marker)
        if not has_section:
            prev.getparent().remove(prev)
        prev = before


def repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    marker = trpr.find(qn("w:tblHeader"))
    if marker is None:
        marker = OxmlElement("w:tblHeader")
        marker.set(qn("w:val"), "true")
        trpr.append(marker)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill="D9E2F3"):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_and_merge_tables(doc):
    tables = list(doc.tables)
    if not tables:
        return
    target = tables[0]
    for source in tables[1:]:
        for row in source.rows[1:]:
            target._tbl.append(deepcopy(row._tr))
        source._tbl.getparent().remove(source._tbl)

    target.alignment = WD_TABLE_ALIGNMENT.CENTER
    target.autofit = False
    widths = [Cm(2.5), Cm(2.1), Cm(0.9), Cm(3.3), Cm(2.0), Cm(3.1), Cm(3.1)]
    for ri, row in enumerate(target.rows):
        if ri == 0:
            repeat_table_header(row)
        for ci, cell in enumerate(row.cells):
            cell.width = widths[ci]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if ri == 0:
                shade_cell(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if (ri == 0 or ci == 2) else WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                set_bidi(paragraph, True)
                for run in paragraph.runs:
                    set_run_fonts(run, size=9.5, bold=(ri == 0))

    caption = OxmlElement("w:p")
    target._tbl.addprevious(caption)
    from docx.text.paragraph import Paragraph

    caption_p = Paragraph(caption, target._parent)
    set_text(caption_p, "جدول ۲-۱: مقایسه پژوهش‌های مرتبط با تشخیص ناهنجاری لاگ", style="Caption", alignment=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_hidden_tc_field(caption_p, "جدول ۲-۱: مقایسه پژوهش‌های مرتبط با تشخیص ناهنجاری لاگ")


def normalize_body(doc):
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style.name in ("Heading 1", "Heading 2", "Heading 3", "عنوان مقدماتی", "برچسب", "Caption"):
            set_bidi(paragraph, True)
            continue
        if text.startswith("Anomaly Detection,"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_bidi(paragraph, False)
            for run in paragraph.runs:
                set_run_fonts(run, size=12)
            continue
        if paragraph.style.name == "Normal" and paragraph.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_after = Pt(0)
        set_bidi(paragraph, True)


def configure_document_settings(doc):
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.5)

    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE, OUTPUT)
    doc = Document(OUTPUT)
    doc.core_properties.title = "پایان‌نامه کارشناسی ارشد مسعود دباغی"
    doc.core_properties.subject = "تشخیص ناهنجاری لاگ‌های BGL با Qwen2.5:7B-Instruct و Ollama"
    doc.core_properties.comments = "نسخه کاری؛ فایل اصلی بدون تغییر حفظ شده است."

    configure_document_settings(doc)
    configure_styles(doc)
    format_front_matter(doc)
    format_headings(doc)
    format_and_merge_tables(doc)
    normalize_body(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
